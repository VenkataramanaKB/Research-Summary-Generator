import requests
from bs4 import BeautifulSoup
import streamlit as st
from docx import Document

st.title('Resume Builder')
st.write('Drop the Google scholar URL and get the summary of the author generated in seconds!')
url =st.text_input(label='Drop the profile url')
 

headers = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/104.0.0.0 Safari/537.36"
}

def write_to_doc(author_results, articles, cited_by):
    # Create a document
    doc = Document()

    # Add the author info
    doc.add_heading('Author Information', level=1)
    doc.add_paragraph(f"Name: {author_results['name']}")
    doc.add_paragraph(f"Position: {author_results['position']}")
    doc.add_paragraph(f"Email: {author_results['email']}")
    doc.add_paragraph(f"Departments: {author_results['departments']}")

    # Add articles
    doc.add_heading('Articles', level=1)
    for article in articles:
        doc.add_paragraph(f"Title: {article['title']}")
        doc.add_paragraph(f"Link: {article['link']}")
        doc.add_paragraph(f"Authors: {article['authors']}")
        doc.add_paragraph(f"Publication: {article['publication']}")
        doc.add_paragraph()  # Add a blank line for spacing

    # Add citation information
    doc.add_heading('Citations', level=1)
    doc.add_paragraph(f"All citations: {cited_by['table'][0]['citations']['all']}")
    doc.add_paragraph(f"Citations since 2017: {cited_by['table'][0]['citations']['since_2017']}")
    doc.add_paragraph(f"All h-index: {cited_by['table'][1]['h_index']['all']}")
    doc.add_paragraph(f"H-index since 2017: {cited_by['table'][1]['h_index']['since_2017']}")
    doc.add_paragraph(f"All i10-index: {cited_by['table'][2]['i_index']['all']}")
    doc.add_paragraph(f"I10-index since 2017: {cited_by['table'][2]['i_index']['since_2017']}")

    # Save the document
    doc_name = 'author_profile.docx'
    doc.save(doc_name)
    return doc_name

if st.button('Search Url') and url:
    response = requests.get(url, headers=headers)
    soup = BeautifulSoup(response.text, 'html.parser')
    author_results = {}
    articles = []
    author_results['name'] = soup.select_one("#gsc_prf_in").text
    author_results['position'] = soup.select_one("#gsc_prf_inw+ .gsc_prf_il").text
    author_results['email'] = soup.select_one("#gsc_prf_ivh").text
    author_results['departments'] = soup.select_one("#gsc_prf_int").text
    for el in soup.select("#gsc_a_b .gsc_a_t"):
        article = {
            'title': el.select_one(".gsc_a_at").text,
            'link': "https://scholar.google.com" + el.select_one(".gsc_a_at")['href'],
            'authors': el.select_one(".gsc_a_at+ .gs_gray").text,
            'publication': el.select_one(".gs_gray+ .gs_gray").text
        }
        articles.append(article)
    for i in range(len(articles)):
        articles[i] = {k: v for k, v in articles[i].items() if v and v != ""}
    cited_by = {}
    cited_by['table'] = []
    cited_by['table'].append({})
    cited_by['table'][0]['citations'] = {}
    cited_by['table'][0]['citations']['all'] = soup.select_one("tr:nth-child(1) .gsc_rsb_sc1+ .gsc_rsb_std").text
    cited_by['table'][0]['citations']['since_2017'] = soup.select_one("tr:nth-child(1) .gsc_rsb_std+ .gsc_rsb_std").text
    cited_by['table'].append({})
    cited_by['table'][1]['h_index'] = {}
    cited_by['table'][1]['h_index']['all'] = soup.select_one("tr:nth-child(2) .gsc_rsb_sc1+ .gsc_rsb_std").text
    cited_by['table'][1]['h_index']['since_2017'] = soup.select_one("tr:nth-child(2) .gsc_rsb_std+ .gsc_rsb_std").text
    cited_by['table'].append({})
    cited_by['table'][2]['i_index'] = {}
    cited_by['table'][2]['i_index']['all'] = soup.select_one("tr~ tr+ tr .gsc_rsb_sc1+ .gsc_rsb_std").text
    cited_by['table'][2]['i_index']['since_2017'] = soup.select_one("tr~ tr+ tr .gsc_rsb_std+ .gsc_rsb_std").text
    st.write(author_results)
    print(type(author_results))
    print(type(articles))
    st.write(articles)
    print(type(cited_by))
    st.write(cited_by['table'])
    doc_name = write_to_doc(author_results, articles, cited_by)
    
    # Provide a download link
    with open(doc_name, "rb") as file:
        st.download_button(label="Download Resume", data=file, file_name=doc_name, mime="application/octet-stream")

